from flask import Flask, render_template, request, send_file
import json, os, markdown
from datetime import datetime
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
import threading
import time
import requests

load_dotenv()

# -----------------------------------
# AI Engine Imports
# -----------------------------------

from ai_engine.writer import generate_blog
from ai_engine.human_variance import reduce_ai_pattern
from ai_engine.detector import get_ai_score
from ai_engine.seo_analyzer import readability_score, seo_score, keyword_density
from ai_engine.topic_suggester import suggest_topics
from ai_engine.ai_seo_analyzer import ai_seo_analysis
from ai_engine.ai_client import generate_analysis_content
from ai_engine.blueprint import generate_blueprint
from ai_engine.human_outline import generate_human_outline
from ai_engine.continuation import continue_blog
from ai_engine.setting import GEMINI_KEYS
from ai_engine.ai_client import BLOCKED_KEYS

# ✅ CHANGE 1: Updated social import
from ai_engine.social_generator import generate_social_posts

# -----------------------------------
# Automation Imports
# -----------------------------------

from automation.image_fetcher import generate_blog_image
from automation.telegram_bot import (
    send_message,
    send_photo,
    send_document,
    send_buttons,
    answer_callback
)

# -----------------------------------
# App Init
# -----------------------------------

app = Flask(__name__)

ANALYTICS_FILE = "analytics_data.json"
USAGE_FILE = "api_usage.json"

# Global state for generated files
generated_word_file = None
generated_image_path = None
generated_instagram_doc = None
generated_linkedin_doc = None

# -----------------------------------
# Constants
# -----------------------------------

PILLARS = [
    "Local SEO",
    "Lead Gen",
    "AI Strategy",
    "Social Media Marketing",
    "Website Optimization"
]

CONTENT_INTENTS = [
    "Educational",
    "Conversion",
    "Authority",
    "SEO",
    "Engagement"
]

# -----------------------------------
# Telegram Session Memory
# -----------------------------------

user_state = {}
user_data = {}


# -----------------------------------
# 🧠 Create Word File with Image
# -----------------------------------

from PIL import Image

def create_word_file(title, blog_content, image_path=None):

    document = Document()
    document.add_heading(title, level=0)

    # ---- Image handling ----
    if image_path and os.path.exists(image_path):
        try:
            img = Image.open(image_path)

            # Convert to PNG for docx compatibility
            png_path = image_path + ".png"
            img.convert("RGB").save(png_path, "PNG")

            document.add_picture(png_path, width=Inches(6))

        except Exception as e:
            print("⚠️ Image conversion failed:", e)

    lines = blog_content.split("\n")

    for line in lines:

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            document.add_heading(line.replace("## ", ""), level=2)

        elif line.startswith("### "):
            document.add_heading(line.replace("### ", ""), level=3)

        else:
            document.add_paragraph(line)

    safe_title = title.replace(" ", "_").replace("/", "").replace(":", "")
    file_path = f"{safe_title}.docx"

    document.save(file_path)

    return file_path


# ✅ CHANGE 2: New social media Word file generator
from PIL import Image

def create_social_word_file(title, post_text, image_path, platform):

    document = Document()

    document.add_heading(f"{platform} Post", level=0)
    document.add_heading(title, level=1)

    # ---- Image handling ----
    if image_path and os.path.exists(image_path):
        try:
            img = Image.open(image_path)

            png_path = image_path + ".png"
            img.convert("RGB").save(png_path, "PNG")

            document.add_picture(png_path, width=Inches(6))

        except Exception as e:
            print("⚠️ Social image conversion failed:", e)

    # Add text
    paragraphs = post_text.split("\n")

    for p in paragraphs:
        p = p.strip()
        if p:
            document.add_paragraph(p)

    safe_title = title.replace(" ", "_").replace("/", "").replace(":", "")

    filename = f"{platform.lower()}_{safe_title}.docx"

    document.save(filename)

    return filename


# -----------------------------------
# 🖼 Async Image Generator
# -----------------------------------

def generate_image_async(topic):
    global generated_image_path
    generated_image_path = generate_blog_image(topic)


# -----------------------------------
# 🧹 Remove AI Garbage Intros
# -----------------------------------

def clean_ai_garbage(text):

    lines = text.split('\n')
    cleaned_lines = []

    garbage_phrases = [
        "here is",
        "here's a",
        "certainly",
        "re-written version",
        "incorporating your",
        "revised article"
    ]

    for line in lines:
        if any(p in line.lower() for p in garbage_phrases) and len(cleaned_lines) < 4:
            continue
        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines).strip()


# -----------------------------------
# 🔬 Semantic Depth Audit
# -----------------------------------

def check_semantic_depth(blog, topic):

    prompt = f"""
Does this blog cover technical details like Schema,
Core Web Vitals, and EEAT for the topic '{topic}'?

Return only:
WEAK
or
STRONG
"""

    result = generate_analysis_content(prompt)

    return result.strip().upper()


# -----------------------------------
# 🧩 Detect Incomplete Blog
# -----------------------------------

def blog_incomplete(blog):

    required_sections = ["faq", "conclusion"]

    for s in required_sections:
        if s not in blog.lower():
            return True

    return False


# -----------------------------------
# 📊 Analytics Persistence
# -----------------------------------

def save_analytics(topic, pillar, intent, read, score):

    entry = {
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "topic": topic,
        "pillar": pillar,
        "intent": intent,
        "readability": read,
        "seo_score": score
    }

    if not os.path.exists(ANALYTICS_FILE):
        with open(ANALYTICS_FILE, "w") as f:
            json.dump([], f)

    with open(ANALYTICS_FILE, "r") as f:
        data = json.load(f)

    data.append(entry)

    with open(ANALYTICS_FILE, "w") as f:
        json.dump(data, f, indent=4)


# -----------------------------------
# 📊 Token Usage Helper
# -----------------------------------

def get_today_token_usage():

    today = datetime.now().strftime("%Y-%m-%d")

    if os.path.exists(USAGE_FILE):

        with open(USAGE_FILE, "r") as f:
            full_usage = json.load(f)

        usage = full_usage.get(today, {"total_tokens": 0})

        return usage.get("total_tokens", 0)

    return 0


# -----------------------------------
# 🏠 Main Web Route
# -----------------------------------

@app.route("/", methods=["GET", "POST"])
def index():

    global generated_word_file
    global generated_image_path
    global generated_instagram_doc
    global generated_linkedin_doc

    if request.method == "POST":

        pillar = request.form.get("pillar")
        custom_topic = request.form.get("custom_topic")
        intent = request.form.get("intent")

        topic = custom_topic.strip() if custom_topic and custom_topic.strip() else pillar

        max_retries = 3
        attempts = 0
        ai_score = 100
        depth_status = "WEAK"

        # -----------------------------------
        # 1️⃣ Human Brainstorm Outline
        # -----------------------------------

        outline = generate_human_outline(topic, pillar)

        # -----------------------------------
        # 2️⃣ Strategic Blueprint
        # -----------------------------------

        blueprint = generate_blueprint(topic, pillar, intent)

        topic_context = f"""
Topic: {topic}

Human Brainstorm Notes:
{outline}

Strategic Blueprint:
{blueprint}

Use these to guide the article.
"""

        # -----------------------------------
        # 3️⃣ Start Async Image Generation
        # -----------------------------------

        image_thread = threading.Thread(target=generate_image_async, args=(topic,))
        image_thread.start()

        # -----------------------------------
        # 4️⃣ Initial Blog Generation
        # -----------------------------------

        blog = generate_blog(topic_context, pillar, intent)
        blog = clean_ai_garbage(blog)

        # -----------------------------------
        # 5️⃣ Autonomous Correction Loop
        # -----------------------------------

        while attempts < max_retries:

            attempts += 1

            ai_score = get_ai_score(blog)
            depth_status = check_semantic_depth(blog, topic)

            if ai_score <= 35 and depth_status == "STRONG":
                break

            paragraphs = blog.split('\n\n')

            if len(paragraphs) > 3:

                intro = paragraphs[0]
                middle = paragraphs[1:-1]
                outro = paragraphs[-1]

                refined_mid = reduce_ai_pattern(
                    '\n\n'.join(middle),
                    retry=True,
                    last_score=ai_score
                )

                blog = f"{intro}\n\n{refined_mid}\n\n{outro}"

            else:

                blog = reduce_ai_pattern(
                    blog,
                    retry=True,
                    last_score=ai_score
                )

            blog = clean_ai_garbage(blog)

        # -----------------------------------
        # 6️⃣ Fix Gemini Truncation
        # -----------------------------------

        if blog_incomplete(blog):
            blog = continue_blog(blog)

        # -----------------------------------
        # ✅ CHANGE 3: 6.5 Generate Social Media Content
        # -----------------------------------

        social_posts = generate_social_posts(topic, blog[:1200])

        print("SOCIAL POSTS:", social_posts)

        instagram_text = ""
        linkedin_text = ""

        if social_posts:

            if "LINKEDIN" in social_posts.upper():

                parts = social_posts.split("LINKEDIN")

                instagram_text = parts[0].replace("INSTAGRAM POST", "").replace("INSTAGRAM", "").strip()

                linkedin_text = parts[-1].strip()

            else:

                instagram_text = social_posts

        # Save TXT files
        with open("instagram_post.txt", "w", encoding="utf-8") as f:
            f.write(instagram_text)

        with open("linkedin_post.txt", "w", encoding="utf-8") as f:
            f.write(linkedin_text)

        # -----------------------------------
        # ✅ CHANGE 4: 7️⃣ Wait for Image Thread & Generate Social Images
        # -----------------------------------

        image_thread.join()
        image_path = generated_image_path

        # Generate social images
        instagram_image = generate_blog_image(topic + " instagram post illustration")
        linkedin_image = generate_blog_image(topic + " linkedin professional graphic")

        # -----------------------------------
        # ✅ CHANGE 5: Export All Word Files
        # -----------------------------------

        # Blog Word File
        generated_word_file = create_word_file(topic, blog, image_path)

        # Instagram Word File
        instagram_doc = create_social_word_file(
            topic,
            instagram_text,
            instagram_image if instagram_image else None,
            "Instagram"
        )
        generated_instagram_doc = instagram_doc

        # LinkedIn Word File
        linkedin_doc = create_social_word_file(
            topic,
            linkedin_text,
            linkedin_image if linkedin_image else None,
            "LinkedIn"
        )
        generated_linkedin_doc = linkedin_doc

        # -----------------------------------
        # 8️⃣ SEO Analysis
        # -----------------------------------

        read = readability_score(blog)
        score = seo_score(blog)
        density = keyword_density(blog)

        analysis = ai_seo_analysis(blog, topic, pillar, intent)

        seo_report = {
            "Readability": read,
            "SEO Score": score,
            "Keyword Density": density,
            "AI_Analysis": analysis,
            "AI_Score": ai_score,
            "Depth_Status": depth_status
        }

        with open("seo_report.json", "w", encoding="utf-8") as f:
            json.dump(seo_report, f, indent=4)

        save_analytics(topic, pillar, intent, read, score)

        return render_template(
            "result.html",
            blog=markdown.markdown(blog),
            instagram_post=instagram_text,
            linkedin_post=linkedin_text,
            instagram_image=instagram_image,
            linkedin_image=linkedin_image,
            readability=read,
            seo_score=score,
            ai_detection_score=ai_score,
            ai_analysis=analysis,
            attempts=attempts,
            image_path=image_path
        )

    return render_template("index.html", pillars=PILLARS)


# -----------------------------------
# 🖼 Image Serving Route
# -----------------------------------

@app.route("/image/<path:filename>")
def serve_image(filename):
    return send_file(filename)


# -----------------------------------
# 📉 Developer Analytics
# -----------------------------------

@app.route("/analytics")
def analytics():

    today_str = datetime.now().strftime("%Y-%m-%d")

    usage_stats = {"total_tokens": 0, "models": {}}

    if os.path.exists(USAGE_FILE):
        with open(USAGE_FILE, "r") as f:
            full_usage = json.load(f)
            usage_stats = full_usage.get(today_str, {"total_tokens": 0, "models": {}})

    blogs_history = []

    if os.path.exists(ANALYTICS_FILE):
        with open(ANALYTICS_FILE, "r") as f:
            blogs_history = json.load(f)

    blogs_today = len([b for b in blogs_history if today_str in b.get("date", "")])
    monthly_total = len(blogs_history)

    avg_score = round(
        sum(d["seo_score"] for d in blogs_history) / len(blogs_history), 2
    ) if blogs_history else 0

    avg_read = round(
        sum(d["readability"] for d in blogs_history) / len(blogs_history), 2
    ) if blogs_history else 0

    tokens_per_blog = round(
        usage_stats.get("total_tokens", 0) / blogs_today
    ) if blogs_today else 0

    active_keys = len(GEMINI_KEYS)

    REQUESTS_PER_KEY = 20
    TOKENS_PER_KEY = 250000

    combined_rpd_limit = REQUESTS_PER_KEY * active_keys
    combined_tpm_limit = TOKENS_PER_KEY * active_keys

    key_health = {}

    for i, key in enumerate(GEMINI_KEYS):
        key_name = f"KEY_{i+1}"
        key_health[key_name] = "Cooling Down" if key in BLOCKED_KEYS else "Healthy"

    return render_template(
        "analytics.html",
        usage=usage_stats,
        total_limit=combined_tpm_limit,
        rpd_limit=combined_rpd_limit,
        active_keys=active_keys,
        key_health=key_health,
        tokens_per_blog=tokens_per_blog,
        data=blogs_history,
        blogs_today=blogs_today,
        monthly_total=monthly_total,
        avg_score=avg_score,
        avg_read=avg_read
    )


# -----------------------------------
# 💡 Topic Suggester Route
# -----------------------------------

@app.route("/suggest-topics", methods=["POST"])
def suggest():

    pillar = request.form.get("pillar")
    intent = request.form.get("intent")
    custom_topic = request.form.get("custom_topic")

    topics = suggest_topics(pillar, intent, custom_topic)

    return render_template(
        "suggestions.html",
        pillar=pillar,
        topics=topics
    )


# -----------------------------------
# 📥 Download Routes
# -----------------------------------

@app.route("/download-blog")
def download_blog():

    global generated_word_file

    if generated_word_file and os.path.exists(generated_word_file):
        return send_file(generated_word_file, as_attachment=True)

    return "File not found", 404


@app.route("/download-seo")
def download_seo():
    return send_file("seo_report.json", as_attachment=True)


# Routes matching HTML buttons exactly (with slash)
@app.route("/download/instagram")
def download_instagram_slash():
    global generated_instagram_doc
    if generated_instagram_doc and os.path.exists(generated_instagram_doc):
        return send_file(generated_instagram_doc, as_attachment=True)
    # Fallback to txt file
    if os.path.exists("instagram_post.txt"):
        return send_file("instagram_post.txt", as_attachment=True)
    return "Instagram file not found", 404


@app.route("/download/linkedin")
def download_linkedin_slash():
    global generated_linkedin_doc
    if generated_linkedin_doc and os.path.exists(generated_linkedin_doc):
        return send_file(generated_linkedin_doc, as_attachment=True)
    # Fallback to txt file
    if os.path.exists("linkedin_post.txt"):
        return send_file("linkedin_post.txt", as_attachment=True)
    return "LinkedIn file not found", 404


# Legacy routes (keep for backward compatibility)
@app.route("/download-instagram")
def download_instagram():
    global generated_instagram_doc
    if generated_instagram_doc and os.path.exists(generated_instagram_doc):
        return send_file(generated_instagram_doc, as_attachment=True)
    return "File not found", 404


@app.route("/download-linkedin")
def download_linkedin():
    global generated_linkedin_doc
    if generated_linkedin_doc and os.path.exists(generated_linkedin_doc):
        return send_file(generated_linkedin_doc, as_attachment=True)
    return "File not found", 404


# -----------------------------------
# 🤖 Telegram Webhook
# -----------------------------------

@app.route("/telegram", methods=["POST"])
def telegram():

    data = request.json
    print("Incoming:", data)

    # -----------------------------------
    # Handle normal messages
    # -----------------------------------

    if "message" in data:

        message = data["message"]
        chat_id = message["chat"]["id"]
        text = message.get("text", "")

        if text == "/start":

            buttons = []

            for p in PILLARS:
                buttons.append([
                    {"text": p, "callback_data": f"pillar:{p}"}
                ])

            send_buttons(
                chat_id,
                "🚀 Welcome to AI Blog SuperAgent\n\nChoose Content Pillar:",
                buttons
            )

            return "ok"

    # -----------------------------------
    # Handle callback buttons
    # -----------------------------------

    if "callback_query" in data:

        query = data["callback_query"]
        chat_id = query["message"]["chat"]["id"]
        callback_data = query["data"]
        callback_id = query["id"]

        answer_callback(callback_id)

        # -----------------------------------
        # Pillar Selected
        # -----------------------------------

        if callback_data.startswith("pillar:"):

            pillar = callback_data.split(":")[1]

            user_data[chat_id] = {"pillar": pillar}

            buttons = [
                [{"text": "Suggest Topics", "callback_data": "topic:suggest"}],
                [{"text": "Custom Topic", "callback_data": "topic:custom"}]
            ]

            send_buttons(chat_id, f"📊 Pillar Selected: {pillar}", buttons)

            return "ok"

        # -----------------------------------
        # Suggest Topics
        # -----------------------------------

        if callback_data == "topic:suggest":

            pillar = user_data[chat_id]["pillar"]

            topics_raw = suggest_topics(pillar, "Educational", None)

            topics_list = topics_raw.split("\n")

            clean_topics = []

            for t in topics_list:

                t = t.strip()

                if not t:
                    continue

                t = t.lstrip("1234567890.- ")

                clean_topics.append(t)

            user_data[chat_id]["topics"] = clean_topics

            buttons = []

            for i in range(min(5, len(clean_topics))):

                buttons.append([
                    {"text": clean_topics[i][:40], "callback_data": f"topic_{i}"}
                ])

            send_buttons(chat_id, "🧠 Choose Topic:", buttons)

            return "ok"

        # -----------------------------------
        # Topic Selected
        # -----------------------------------

        if callback_data.startswith("topic_"):

            index = int(callback_data.split("_")[1])

            topic = user_data[chat_id]["topics"][index]

            user_data[chat_id]["topic"] = topic

            buttons = []

            for intent in CONTENT_INTENTS:

                buttons.append([
                    {"text": intent, "callback_data": f"intent:{intent}"}
                ])

            send_buttons(chat_id, f"📝 Topic Selected:\n\n{topic}\n\nChoose Intent:", buttons)

            return "ok"

        # -----------------------------------
        # Intent Selected
        # -----------------------------------

        if callback_data.startswith("intent:"):

            intent = callback_data.split(":")[1]

            user_data[chat_id]["intent"] = intent

            buttons = [[
                {"text": "🚀 Generate Blog", "callback_data": "generate_blog"}
            ]]

            send_buttons(chat_id, f"Intent Selected: {intent}", buttons)

            return "ok"

        # -----------------------------------
        # Generate Blog + Social Content
        # -----------------------------------

        if callback_data == "generate_blog":

            topic = user_data[chat_id]["topic"]
            pillar = user_data[chat_id]["pillar"]
            intent = user_data[chat_id]["intent"]

            send_message(chat_id, "🧠 Generating blog... Please wait (~20 seconds)")

            try:

                # -----------------------------------
                # Generate Blog
                # -----------------------------------

                blog_content = generate_blog(topic, pillar, intent)

                blog_content = clean_ai_garbage(blog_content)

                if blog_incomplete(blog_content):
                    blog_content = continue_blog(blog_content)

                # -----------------------------------
                # Blog Image
                # -----------------------------------

                image_path = generate_blog_image(topic)

                if image_path and os.path.exists(image_path):
                    send_photo(chat_id, image_path)

                # -----------------------------------
                # Generate Social Posts
                # -----------------------------------

                social_posts = generate_social_posts(topic, blog_content[:1200])

                instagram_text = ""
                linkedin_text = ""

                if "LINKEDIN POST" in social_posts:

                    parts = social_posts.split("LINKEDIN POST")

                    instagram_text = parts[0].replace("INSTAGRAM POST", "").strip()
                    linkedin_text = parts[1].strip()

                else:
                    instagram_text = social_posts

                # -----------------------------------
                # Social Images
                # -----------------------------------

                instagram_image = generate_blog_image(topic + " instagram post illustration")
                linkedin_image = generate_blog_image(topic + " linkedin professional graphic")

                # -----------------------------------
                # Send Instagram Post
                # -----------------------------------

                send_message(chat_id, f"📸 Instagram Post:\n\n{instagram_text}")

                if instagram_image and os.path.exists(instagram_image):
                    send_photo(chat_id, instagram_image)

                # -----------------------------------
                # Send LinkedIn Post
                # -----------------------------------

                send_message(chat_id, f"💼 LinkedIn Post:\n\n{linkedin_text}")

                if linkedin_image and os.path.exists(linkedin_image):
                    send_photo(chat_id, linkedin_image)

                # -----------------------------------
                # Blog Preview
                # -----------------------------------

                preview = blog_content[:500]

                send_message(
                    chat_id,
                    f"""
📝 Blog Generated

Title: {topic}
Pillar: {pillar}
Intent: {intent}

Preview 👇

{preview}...
"""
                )

                # -----------------------------------
                # Create Word Files
                # -----------------------------------

                blog_doc = create_word_file(topic, blog_content, image_path)

                instagram_doc = create_social_word_file(
                    topic,
                    instagram_text,
                    instagram_image,
                    "Instagram"
                )

                linkedin_doc = create_social_word_file(
                    topic,
                    linkedin_text,
                    linkedin_image,
                    "LinkedIn"
                )

                # -----------------------------------
                # Send Word Files
                # -----------------------------------

                send_document(chat_id, blog_doc)
                send_document(chat_id, instagram_doc)
                send_document(chat_id, linkedin_doc)

            except Exception as e:

                print("Telegram blog generation error:", e)

                send_message(chat_id, "⚠️ Error generating blog. Please try again.")

            return "ok"

    return "ok"

# -----------------------------------
# 🔁 Keep Server Alive
# -----------------------------------

def keep_alive():

    while True:
        try:
            requests.get("https://shivam-blog-superagent-7.onrender.com")
            print("Server pinged")
        except Exception as e:
            print("Ping failed:", e)

        time.sleep(600)


threading.Thread(target=keep_alive, daemon=True).start()


# -----------------------------------
# 🚀 Run Server
# -----------------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)